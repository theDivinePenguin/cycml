"""Generate a beautiful, publication-quality Word Document (.docx) for START_TO_FINISH_PROJECT_REPORT.md."""
import os
from pathlib import Path
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn


def set_cell_background(cell, fill_hex):
    """Set shading/background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set cell padding in dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def add_styled_heading(doc, text, level):
    """Add styled headings with consistent corporate colors and typography."""
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Arial"
    run.bold = True
    
    if level == 1:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(16, 54, 110)   # Deep Navy
        
        # Add bottom accent line
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="4" w:color="0284C7"/></w:pBdr>')
        pPr.append(pBdr)
    elif level == 2:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(3, 105, 161)   # Blue
    elif level == 3:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(51, 65, 85)    # Slate
    return p


def add_bullet(doc, text, bold_prefix="", indent_level=0):
    """Add a clean bullet with dark text and custom spacing."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    if indent_level > 0:
        p.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
    
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = "Arial"
        r_pre.font.size = Pt(10)
        r_pre.font.bold = True
        r_pre.font.color.rgb = RGBColor(15, 23, 42)
    
    r_body = p.add_run(text)
    r_body.font.name = "Arial"
    r_body.font.size = Pt(10)
    r_body.font.color.rgb = RGBColor(33, 37, 41)
    return p


def add_body_paragraph(doc, text, bold_prefix=""):
    """Add standard body paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = "Arial"
        r_pre.font.size = Pt(10)
        r_pre.font.bold = True
        r_pre.font.color.rgb = RGBColor(16, 54, 110)
        
    r_body = p.add_run(text)
    r_body.font.name = "Arial"
    r_body.font.size = Pt(10)
    r_body.font.color.rgb = RGBColor(33, 37, 41)
    return p


def add_callout_box(doc, text, bold_title=""):
    """Add a highlighted callout box with a colored left border."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.8)
    set_cell_background(cell, "F0F9FF")  # Light blue background
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    # Left border only (accent cyan)
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="0284C7"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_title:
        r_title = p.add_run(bold_title + "\n")
        r_title.font.name = "Arial"
        r_title.font.size = Pt(10.5)
        r_title.font.bold = True
        r_title.font.color.rgb = RGBColor(3, 105, 161)
        
    r_body = p.add_run(text)
    r_body.font.name = "Arial"
    r_body.font.size = Pt(9.5)
    r_body.font.color.rgb = RGBColor(15, 23, 42)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_image_with_caption(doc, img_path, caption, width_in=6.4):
    """Add centered image with styled caption below."""
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(2)
    p_img.paragraph_format.keep_with_next = True
    
    if Path(img_path).exists():
        p_img.add_run().add_picture(str(img_path), width=Inches(width_in))
    else:
        p_img.add_run(f"[Image Missing: {img_path}]").font.color.rgb = RGBColor(220, 38, 38)
        
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(2)
    p_cap.paragraph_format.space_after = Pt(10)
    r_cap = p_cap.add_run(caption)
    r_cap.font.name = "Arial"
    r_cap.font.size = Pt(9)
    r_cap.font.bold = True
    r_cap.font.color.rgb = RGBColor(100, 116, 139)


def generate_docx():
    doc = docx.Document()
    
    # 1. Page Margins (0.8 inches all around)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
        # Header & Footer
        footer = section.footer
        p_ft = footer.paragraphs[0]
        p_ft.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_ft = p_ft.add_run("DeepCycloNet | Smart India Hackathon 2026 — PS ID 26070")
        r_ft.font.name = "Arial"
        r_ft.font.size = Pt(8.5)
        r_ft.font.color.rgb = RGBColor(148, 163, 184)

    # 2. Document Title Banner
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(4)
    p_title.paragraph_format.space_after = Pt(2)
    r_main = p_title.add_run("DeepCycloNet: Start-to-Finish Project Report")
    r_main.font.name = "Arial"
    r_main.font.size = Pt(22)
    r_main.font.bold = True
    r_main.font.color.rgb = RGBColor(16, 54, 110)
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(12)
    r_sub = p_sub.add_run("AI/ML System for Tropical Cyclone Tracking, Intensity Forecasting & Rapid Intensification Early Warning\n")
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(12)
    r_sub.font.bold = True
    r_sub.font.color.rgb = RGBColor(3, 105, 161)
    
    r_meta = p_sub.add_run("Smart India Hackathon (SIH) 2026 — Problem Statement ID 26070\n"
                           "Code Repository: https://github.com/theDivinePenguin/cycml | Live Workstation: https://thedivinepenguin.github.io/cycml/")
    r_meta.font.name = "Arial"
    r_meta.font.size = Pt(9.5)
    r_meta.font.color.rgb = RGBColor(100, 116, 139)
    
    # ---------------------------------------------------------
    # Executive Summary
    # ---------------------------------------------------------
    add_styled_heading(doc, "Executive Summary: What We Built in 60 Seconds", 1)
    
    add_body_paragraph(doc, 
        "When a tropical cyclone forms over the ocean, the single most dangerous threat to human life is Rapid Intensification (RI) "
        "— when an ordinary, moderate storm unexpectedly explodes into a catastrophic super cyclone within 24 hours right before slamming into the coast. "
        "Traditional meteorological methods frequently miss this sudden acceleration or issue warnings too late, leaving coastal disaster authorities "
        "with insufficient lead time to evacuate millions of vulnerable citizens."
    )
    
    add_callout_box(doc,
        "• 18–21 Hour Advance Warning: Accurately anticipates explosive rapid intensification up to a full day before peak destructive landfall.\n"
        "• 40% to 45% Lower Errors: Delivers continuous wind speed forecasts with 4.98 knots error at +6h and 10.75 knots at +24h, outperforming official IMD agency guidelines (~8.2 kt and ~18.5 kt).\n"
        "• Dual-Brain Architecture: Simultaneously analyzes an 18-hour historical sequence of satellite cloud images (the 'Eye in the Sky') and oceanic heat energy measurements (the 'Fuel Tank Below').\n"
        "• Operational Workstation: Deployed as a mission-control grade interactive dashboard with multi-spectral satellite viewing and live hazard telemetry.",
        "KEY INNOVATION BREAKTHROUGHS:"
    )

    # ---------------------------------------------------------
    # Chapter 1: The Problem
    # ---------------------------------------------------------
    add_styled_heading(doc, "Chapter 1: The Deadly Cyclone Problem & Operational Blindspot", 1)
    
    add_body_paragraph(doc,
        "Every cyclone season, storms forming in the Bay of Bengal and Arabian Sea threaten millions of lives in Odisha, Andhra Pradesh, "
        "Tamil Nadu, West Bengal, and Gujarat. While track path forecasting has improved over the past decades, intensity forecasting — particularly "
        "predicting when a storm will rapidly strengthen — has remained an unresolved grand challenge in atmospheric science."
    )
    
    add_styled_heading(doc, "Why Traditional Meteorological Forecasting Struggles:", 2)
    add_bullet(doc, "Looking at a single satellite snapshot is like guessing where a race car will finish by looking at a picture of it parked. A cyclone is a dynamic vortex; knowing its current direction requires seeing its rotation momentum over time.", "1. The Single-Snapshot Trap: ")
    add_bullet(doc, "A cyclone is powered from below. If ocean water is blistering hot (>28.5°C) with deep thermal content, even a loose cloud cluster can rapidly intensify overnight. Models that ignore ocean data miss the fuel powering the storm.", "2. The Ocean Heat Blindspot: ")
    add_bullet(doc, "When forecasters extrapolate current wind speeds 24 hours ahead, they frequently predict peak storm intensity after the cyclone has already made landfall and begun dying over land, issuing false alarms when danger has passed.", "3. The 24-Hour Persistence Lag: ")
    add_bullet(doc, "Mass coastal evacuations cost state exchequers ₹50 to ₹100+ Crores per event. Ordering unnecessary evacuations for storms dying over cooler waters drains emergency budgets and breeds public cynicism.", "4. Massive Economic Evacuation Costs: ")

    # ---------------------------------------------------------
    # Chapter 2: The Data Journey
    # ---------------------------------------------------------
    add_styled_heading(doc, "Chapter 2: The Data Journey (Building the Ground Truth)", 1)
    
    add_body_paragraph(doc,
        "To ensure true operational reliability across global cyclone basins, we trained and evaluated DeepCycloNet on the comprehensive "
        "Tropical Cyclone Image Dataset (TCIR), spanning 70,499 satellite images across 1,285 unique tropical cyclones over 15+ continuous years (1998–2017)."
    )
    
    add_styled_heading(doc, "The 3 Satellite Modalities Processed 24/7:", 2)
    add_bullet(doc, "Measures cloud-top brightness temperatures in physical Kelvin. Colder cloud tops (-75°C to -85°C) identify violent eyewall thunderstorm updrafts. Operates 24/7 day and night.", "1. Clean Infrared (IR1 10.8 µm): ")
    add_bullet(doc, "Measures mid-to-upper tropospheric moisture, revealing dry air intrusions that can choke cyclone development.", "2. Water Vapor (WV 6.7 µm): ")
    add_bullet(doc, "High-resolution optical sunlight reflection showing fine cloud textures, rainbands, and eyewall symmetry. Our system includes learned day-night gating that automatically deactivates optical channels at sunset.", "3. Visible Albedo (VIS 0.65 µm): ")
    
    add_styled_heading(doc, "The 5 Atmospheric & Oceanic Parameters (SHIPS):", 2)
    add_bullet(doc, "Warm water temperature powering the storm (cyclones need ≥ 26.5°C to develop).", "• Sea Surface Temperature (SST): ")
    add_bullet(doc, "Measures the depth and reservoir of warm water below the surface — the cyclone's true fuel tank.", "• Ocean Heat Content (OHC): ")
    add_bullet(doc, "Wind speed differences between upper and lower atmospheric layers. High shear tears storms apart; low shear allows explosive growth.", "• Vertical Wind Shear: ")
    add_bullet(doc, "Moisture at 3–5 km altitude. Dry air prevents cloud clusters from organizing into a coherent eye.", "• Mid-Level Relative Humidity: ")
    add_bullet(doc, "Central atmospheric surface pressure drop driving peak wind speeds.", "• Minimum Sea-Level Pressure (MSLP): ")

    # ---------------------------------------------------------
    # Chapter 3: The DeepCycloNet Architecture
    # ---------------------------------------------------------
    add_styled_heading(doc, "Chapter 3: The DeepCycloNet Solution (How the AI Works)", 1)
    
    add_body_paragraph(doc,
        "DeepCycloNet combines spatial computer vision, historical time-series memory, and oceanic physics into a unified multi-modal neural architecture."
    )
    
    add_image_with_caption(doc, "figures/slide3_technical_architecture_flowchart.png", 
                           "Figure 1: DeepCycloNet End-to-End Multi-Modal AI Architecture Flowchart (Inputs → Fusion Transformer → 3 Operational Heads).")

    add_styled_heading(doc, "The 3 Integrated AI Processing Stages:", 2)
    add_bullet(doc, "A modified ResNet-18 neural network processes each satellite frame, identifying the eye core, the spiral rainband radius, and the symmetry of convective cloud tops.", "Stage 1 — Spatial Computer Vision (ResNet-18): ")
    add_bullet(doc, "An 18-hour sequence of 7 consecutive frames (t-18h, t-15h, t-12h, t-9h, t-6h, t-3h, t) is fed through an 8-head Temporal Transformer. This allows the AI to measure rotational momentum and storm evolution over time.", "Stage 2 — 18-Hour Spatio-Temporal Transformer: ")
    add_bullet(doc, "A multi-head cross-attention layer fuses visual cloud representations with the 5 oceanic parameters. The model cross-examines: 'Are eyewall updrafts accelerating, AND is ocean heat content sufficient to sustain explosive intensification?'", "Stage 3 — Cross-Attention Multi-Modal Fusion: ")

    # ---------------------------------------------------------
    # Chapter 4: The Three Operational Outputs
    # ---------------------------------------------------------
    add_styled_heading(doc, "Chapter 4: The Three Operational Forecast Outputs", 1)
    
    add_body_paragraph(doc,
        "Rather than outputting a single ambiguous number, DeepCycloNet provides three synchronized operational outputs designed for actionable decision-making:"
    )
    
    add_styled_heading(doc, "1. Rapid Intensification (RI) Hazard Early Warning:", 2)
    add_body_paragraph(doc,
        "Outputs the exact mathematical probability P(RI) that the cyclone's wind speed will surge by 30 knots or more in the next 24 hours. "
        "When P(RI) crosses our operational decision threshold (τ = 0.016), an automated alert is triggered 18 to 21 hours before peak destructive winds develop."
    )
    
    add_styled_heading(doc, "2. 24-Hour Macro Trend Classification:", 2)
    add_body_paragraph(doc,
        "Classifies the storm's 24-hour evolutionary regime into WEAKENING (decaying by ≥ 10 kt), STABLE (within ±10 kt), or INTENSIFYING (growing by ≥ 10 kt). "
        "The model achieves 64.71% overall accuracy and a critical 79.0% recall on weakening events, giving disaster authorities confidence when landfalls are calming down."
    )
    
    add_styled_heading(doc, "3. Continuous Wind Speed Forecasts (+6h, +12h, +24h):", 2)
    add_body_paragraph(doc,
        "Delivers exact maximum sustained wind speeds in knots at 6-hour, 12-hour, and 24-hour horizons. Predictions are mapped directly into official "
        "IMD storm categories (Cyclonic Storm, Severe, Very Severe, Extremely Severe, Super Cyclone) and Saffir-Simpson categories (Cat 1 to Cat 5)."
    )

    # ---------------------------------------------------------
    # Chapter 5: Real-World Testing & Verification
    # ---------------------------------------------------------
    add_styled_heading(doc, "Chapter 5: Real-World Testing & Verification (The Proving Ground)", 1)
    
    add_body_paragraph(doc,
        "To verify that our model does not simply memorize training samples, we tested it against 7,901 held-out sequences from 187 completely unseen cyclones "
        "with zero data overlap. The system proved highly accurate across diverse storm types, peak intensities, and ocean basins."
    )
    
    add_image_with_caption(doc, "figures/slide4_least_error_cyclones.png",
                           "Figure 2: Real-World AI Performance on 2 Major Severe Cyclones with Least Error (Typhoon Guchol & Hurricane Blas).")

    add_styled_heading(doc, "Showcase Storm 1 — Typhoon Guchol (West Pacific, Peak: 105 kt / Category 4):", 2)
    add_body_paragraph(doc,
        "Across 35 consecutive operational cycles spanning over 100 hours of storm lifecycle, DeepCycloNet achieved a remarkable +24h Mean Absolute Error "
        "of just 5.46 knots and a 94.3% trend classification accuracy. The AI anticipated both the rapid intensification ramp to 105 kt and the post-peak weakening phase."
    )
    
    add_styled_heading(doc, "Showcase Storm 2 — Hurricane Blas (East Pacific, Peak: 120 kt / Category 4):", 2)
    add_body_paragraph(doc,
        "Tracking Blas as it developed from a 65 kt tropical storm into a 120 kt major hurricane, DeepCycloNet achieved a +24h MAE of 6.63 knots "
        "and an 83.1% trend accuracy, accurately capturing the peak eyewall strength without over-prediction or delay."
    )

    add_styled_heading(doc, "Showcase Storm 3 — Super Cyclone Phet (Arabian Sea, Peak: 125 kt / Category 4):", 2)
    add_body_paragraph(doc,
        "In May 2010, Super Cyclone Phet threatened coastal Oman and Pakistan. DeepCycloNet's RI early warning engine surged past the critical hazard threshold "
        "18 hours before Phet reached its Category 4 peak, providing crucial lead time before severe coastal impacts occurred."
    )

    # ---------------------------------------------------------
    # Chapter 6: Head-to-Head Comparison vs Weather Agencies
    # ---------------------------------------------------------
    add_styled_heading(doc, "Chapter 6: Head-to-Head Benchmark vs Weather Agencies", 1)
    
    add_body_paragraph(doc,
        "We compared DeepCycloNet's continuous wind speed forecast errors against authoritative operational benchmarks from the India Meteorological Department (IMD) "
        "and the Joint Typhoon Warning Center (JTWC)."
    )
    
    add_image_with_caption(doc, "figures/slide6_benchmark_comparison.png",
                           "Figure 3: Authoritative Performance Benchmark vs Official Weather Agency Errors (39% to 45% Error Reductions).")

    # Table of Benchmark Comparison
    table = doc.add_table(rows=5, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ["Forecast Lead Time", "Traditional Agency Error (IMD / JTWC)", "DeepCycloNet AI Error", "Error Reduction / Gain"]
    for col_idx, h_text in enumerate(headers):
        cell = table.cell(0, col_idx)
        set_cell_background(cell, "10366E")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h_text)
        r.font.name = "Arial"
        r.font.size = Pt(9.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)

    rows_data = [
        ("+6 Hours Ahead", "8.2 knots", "4.98 knots", "-39.3% Error Reduction"),
        ("+12 Hours Ahead", "12.8 knots", "6.99 knots", "-45.4% Error Reduction"),
        ("+24 Hours Ahead", "18.5 knots", "10.75 knots", "-41.9% Error Reduction"),
        ("Rapid Intensification Warning", "0 to 6 Hours Notice", "18 to 21 Hours Notice", "+12 to 15 Hours Extra Lead Time"),
    ]

    for row_idx, r_data in enumerate(rows_data):
        fill_bg = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, val in enumerate(r_data):
            cell = table.cell(row_idx + 1, col_idx)
            set_cell_background(cell, fill_bg)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            r.font.name = "Arial"
            r.font.size = Pt(9)
            if col_idx == 2 or col_idx == 3:
                r.font.bold = True
                r.font.color.rgb = RGBColor(3, 105, 161) if col_idx == 2 else RGBColor(4, 120, 87)
            else:
                r.font.color.rgb = RGBColor(33, 37, 41)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ---------------------------------------------------------
    # Chapter 7: The Live Workstation
    # ---------------------------------------------------------
    add_styled_heading(doc, "Chapter 7: The Live Operational Meteorological Workstation", 1)
    
    add_body_paragraph(doc,
        "Rather than leaving our models in research notebooks, we engineered and deployed a serious, mission-control grade Meteorological Analysis Workstation "
        "designed specifically for operational duty officers and disaster management centers."
    )
    
    add_image_with_caption(doc, "figures/workstation_preview.png",
                           "Figure 4: The Live Deployed Meteorological Analysis Workstation (No AI Slop — Professional Mission-Control Density).")

    add_styled_heading(doc, "Key Workstation Capabilities:", 2)
    add_bullet(doc, "Allows instant switching between Clean Infrared (calibrated Dvorak Kelvin scale 190–310 K), Water Vapor (6.7 µm), Visible (0.65 µm), and neural network Cross-Attention Saliency maps.", "• Multi-Band Satellite Observation Deck: ")
    add_bullet(doc, "Provides concentric range rings at 100 km, 200 km, and 300 km from the storm center with a central reticle.", "• Eyewall Range Rings & Reticle: ")
    add_bullet(doc, "Forecasters can scrub backwards and forwards through the 18-hour storm history (t-18h to NOW) using a live DVR player at 1X, 2X, or 4X speed.", "• 7-Frame Sequence Strip & DVR Scrubber: ")
    add_bullet(doc, "Live horizontal risk gauge marking the operational decision line (τ = 0.016) with immediate warning status badges.", "• RI Hazard Risk Gauge: ")
    add_bullet(doc, "Dual-line interactive charts comparing AI predictions against ground-truth best track records, featuring an EMA presentation smoothing toggle while preserving raw sensor telemetry.", "• Forecast Proving Ground: ")

    # ---------------------------------------------------------
    # Chapter 8: Project Roadmap & Real-World Impact
    # ---------------------------------------------------------
    add_styled_heading(doc, "Chapter 8: Project Roadmap & Real-World Impact", 1)
    
    add_image_with_caption(doc, "figures/slide5_roadmap_infographic.png",
                           "Figure 5: Project Evolution Roadmap (What Was Done, What We Did, and What We Could Do).")

    add_styled_heading(doc, "1. What Was Done (Baseline):", 2)
    add_body_paragraph(doc, "Single-frame static image models; satellite-only architectures ignoring ocean thermodynamics; continuous regression prone to severe 24-hour persistence lag.")
    
    add_styled_heading(doc, "2. What We Did (DeepCycloNet Innovation):", 2)
    add_body_paragraph(doc, "18-hour Spatio-Temporal Transformer (K=7); multi-modal ocean fusion (SHIPS SST, OHC, Shear, MSLP, RH); 3-task cost-sensitive model; live web workstation.")
    
    add_styled_heading(doc, "3. What We Could Do (Future Operational Scope):", 2)
    add_bullet(doc, "Establish automated API ingestion pipelines directly connecting India's INSAT-3D/3DR geostationary satellites for real-time coverage over the Bay of Bengal and Arabian Sea.", "• Direct INSAT-3D / 3DR Satellite Integration: ")
    add_bullet(doc, "Extend the temporal transformer to simultaneously predict latitude/longitude coordinates and cone-of-uncertainty landfall locations alongside wind speed.", "• Track Path & Landfall Cone Forecasting: ")
    add_bullet(doc, "Hook the RI early warning trigger directly into automated SMS and mobile dispatch networks for coastal District Magistrates and NDRF battalions.", "• Automated Emergency Warning Dispatches: ")

    # ---------------------------------------------------------
    # Conclusion
    # ---------------------------------------------------------
    add_styled_heading(doc, "Conclusion: Impact on India (SIH 26070)", 1)
    
    add_body_paragraph(doc,
        "DeepCycloNet directly answers Smart India Hackathon Problem Statement 26070 by delivering an end-to-end, scientifically validated, "
        "and operationally deployed tropical cyclone intelligence system."
    )
    
    add_callout_box(doc,
        "1. Economic Savings (₹50–100+ Crores per Event): Eliminates unnecessary coastal evacuations when storms are decaying while guaranteeing full mobilization when danger is severe.\n"
        "2. Social Benefit (Zero Human Casualties): Direct early warning support for vulnerable coastal populations in Odisha, Andhra Pradesh, Tamil Nadu, and Gujarat.\n"
        "3. Technological Sovereignty: Fully reproducible, open-source AI architecture validated on 15+ years of real-world storm data.",
        "THE LIFESAVING IMPACT AT A GLANCE:"
    )

    out_docx_path = Path("reports/START_TO_FINISH_PROJECT_REPORT.docx")
    out_docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_docx_path)
    print(f"Successfully generated DOCX report: {out_docx_path} ({out_docx_path.stat().st_size / 1024 / 1024:.2f} MB)")


if __name__ == "__main__":
    generate_docx()
